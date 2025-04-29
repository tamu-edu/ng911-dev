import os
import xml.etree.ElementTree as ET
import json
import csv
from docx import Document
from docx.shared import Pt
from junit_xml import TestSuite, TestCase

from .format_types.PDFService import PDFService
from .report_enums import ReportType
from services.config.schemas.requirements_schema import REQUIREMENTS_SCHEMA


class ReportService:
    report_data: dict = {}

    def __init__(self, test_oracle):
        """Initialize with a TestOracle object."""
        self.test_oracle = test_oracle
        self.__generate_report_data()

    def generate_report(self, report_format: str, output_path: str):
        """Generate a report in the specified format."""
        if report_format.lower() == ReportType.PDF.value:
            self.__generate_pdf(output_path)
        elif report_format.lower() == ReportType.DOCUMENT.value:
            self.__generate_docx(output_path)
        elif report_format.lower() == ReportType.XML.value:
            self.__generate_xml(output_path)
        elif report_format.lower() == ReportType.JSON.value:
            self.__generate_json(output_path)
        elif report_format.lower() == ReportType.CSV.value:
            self.__generate_csv(output_path)
        else:
            raise ValueError(f"Unsupported report format. Use one of {ReportType.list()}")

    def __get_var_description(self, var_name: str) -> str:
        test_config = self.test_oracle.get_test_config()
        for test in test_config.conformance.tests:
            for var in test.variations:
                if var.name == var_name:
                    return var.description

    def _get_all_variations_names(self, variations: list):

        return {
            "names": [var.name for var in variations],
            "descriptions": {
                var.name: self.__get_var_description(var.name) for var in variations
            }
        }

    def _generate_req_vars_results(self, variation_names: list, req, var_descriptions):
        req_dict = {}
        req_schema = REQUIREMENTS_SCHEMA.get(req.name)
        req_dict["req_text"] = req_schema["requirement_text"]
        req_dict["section"] = req_schema["document_section"]
        req_dict["description"] = req_schema["description"]
        for var_name in variation_names:
            variation_dict = {
                "description": var_descriptions[var_name]
            }

            for scenario in self.test_oracle.scenarios:
                if var_name == scenario.name:
                    subtests = req_schema.get("subtests", [])

                    for verdict in scenario.intermediate_verdicts:
                        if verdict.test_name in subtests or len(subtests) == 0:
                            variation_dict[verdict.test_name] = verdict.test_verdict.value

            req_dict[var_name] = variation_dict

        return req_dict

    def __generate_report_data(self):
        run_config = self.test_oracle.get_run_config()

        for test in run_config.tests:
            test_dict = {}
            all_variations_names = self._get_all_variations_names(test.variations)

            for req in test.requirements:
                if "all" in req.variations:
                    test_dict[req.name] = self._generate_req_vars_results(
                        all_variations_names.get("names"), req, all_variations_names.get("descriptions")
                    )
                else:
                    test_dict[req.name] = self._generate_req_vars_results(
                        req.variations, req, all_variations_names.get("descriptions")
                    )

            self.report_data[test.name] = test_dict

    def __generate_pdf(self, output_path):
        """Generate a high-quality PDF report."""

        pdf = PDFService()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        page_width = pdf.w - 2 * pdf.l_margin

        for test_name, requirements in self.report_data.items():
            pdf.set_font("Arial", "B", 12)
            pdf.cell(page_width, 10, f"Test Case: {test_name}", ln=True)
            pdf.ln(2)

            for req_name, req_data in requirements.items():
                pdf.set_font("Arial", "B", 11)
                pdf.set_text_color(0, 51, 102)
                pdf.cell(page_width, 8, f"Requirement: {req_name}", ln=True)

                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Arial", "", 10)
                pdf.safe_multi_cell(page_width, 6, f"Description: {req_data.get('description', '')}")
                pdf.ln(1)

                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Arial", "", 10)
                req_text = f"Requirement Text: {req_data.get('req_text', '')}"
                pdf.safe_multi_cell(page_width, 6, req_text)
                pdf.ln(3)

                for var_name in [v for v in req_data if v not in ['description', 'req_text', 'section']]:
                    var_info = req_data[var_name]

                    pdf.set_font("Arial", "I", 10)
                    pdf.cell(page_width, 6, f"Variation: {var_name}", ln=True)
                    pdf.set_font("Arial", "", 10)
                    pdf.safe_multi_cell(page_width, 6, f"Variation Description: {var_info.get('description', '')}")
                    pdf.ln(1)

                    # Table header
                    col1_width = page_width * 0.7
                    col2_width = page_width * 0.3

                    pdf.set_font("Arial", "B", 10)
                    pdf.set_fill_color(220, 220, 220)
                    pdf.cell(col1_width, 8, "Check", border=1, fill=True)
                    pdf.cell(col2_width, 8, "Verdict", border=1, fill=True, ln=True)

                    # Table rows
                    pdf.set_font("Arial", "", 10)
                    for check_name, verdict in var_info.items():
                        if check_name == "description":
                            continue
                        pdf.cell(col1_width, 8, pdf.sanitize_text(check_name), border=1)
                        pdf.cell(col2_width, 8, pdf.sanitize_text(verdict), border=1, ln=True)

                    pdf.ln(6)

                pdf.ln(8)

            pdf.ln(10)

        pdf.output(output_path)

    def __generate_docx(self, output_path):
        """Generate a DOCX report with one table per test case."""
        doc = Document()
        doc.add_heading("Conformance Test Report", level=1)

        for test_name, requirements in self.report_data.items():
            doc.add_heading(f"Test Case: {test_name}", level=2)

            for req_name, req_data in requirements.items():
                doc.add_heading(f"Requirement: {req_name}", level=3)
                doc.add_paragraph(f"Description: {req_data['description']}")
                doc.add_paragraph(f"Text: {req_data['req_text']}")
                doc.add_paragraph(f"Section: {req_data['section']}")

                for var_name in [v for v in req_data if v not in ['description', 'req_text', 'section']]:
                    var_info = req_data[var_name]
                    doc.add_heading(f"Variation: {var_name}", level=4)
                    doc.add_paragraph(f"Description: {var_info['description']}")

                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Check"
                    hdr_cells[1].text = "Verdict"

                    for check_name, verdict in var_info.items():
                        if check_name == "description":
                            continue
                        row_cells = table.add_row().cells
                        row_cells[0].text = check_name
                        row_cells[1].text = verdict

                        for paragraph in row_cells[0].paragraphs + row_cells[1].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                    doc.add_paragraph()

        doc.save(output_path)

    def __generate_xml(self, output_path):
        """Generate an XML (JUnitXML) report."""
        testsuites = ET.Element("testsuites")

        for test_name, requirements in self.report_data.items():
            testsuite = ET.SubElement(testsuites, "testsuite", name=test_name)

            for req_name, req_data in requirements.items():
                for var_name in [v for v in req_data if v not in ['description', 'req_text', 'section']]:
                    testcase = ET.SubElement(testsuite, "testcase", classname=req_name, name=var_name)
                    var_info = req_data[var_name]

                    for check_name, verdict in var_info.items():
                        if check_name == "description":
                            continue
                        if verdict.upper() != "PASSED":
                            failure = ET.SubElement(testcase, "failure", message=f"{check_name}: {verdict}")
                            failure.text = f"{check_name} - {verdict}"

        tree = ET.ElementTree(testsuites)
        ET.indent(tree, space="  ", level=0)
        tree.write(output_path, encoding="utf-8", xml_declaration=True)

    def __generate_json(self, output_path):
        """Generate a JSON report."""
        with open(output_path, "w") as json_file:
            json.dump(self.report_data, json_file, indent=4)

    def __generate_csv(self, output_path):
        """Generate a CSV report."""
        with open(output_path, mode="w", newline="") as csv_file:
            writer = csv.writer(csv_file)
            writer.writerow(
                ["Test Case", "Requirement", "Requirement Description", "Requirement Text", "Section", "Variation",
                 "Variation Description", "Check", "Verdict"])

            for test_name, requirements in self.report_data.items():
                for req_name, req_data in requirements.items():
                    req_desc = req_data["description"]
                    req_text = req_data["req_text"]
                    req_section = req_data["section"]

                    for var_name in [v for v in req_data if v not in ['description', 'req_text', 'section']]:
                        var_info = req_data[var_name]
                        var_desc = var_info.get("description", "")

                        for check_name, verdict in var_info.items():
                            if check_name == "description":
                                continue

                            writer.writerow([
                                test_name,
                                req_name,
                                req_desc,
                                req_text,
                                req_section,
                                var_name,
                                var_desc,
                                check_name,
                                verdict
                            ])
